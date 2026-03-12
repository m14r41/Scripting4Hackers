import os
import fitz  # PyMuPDF
from PIL import Image
from io import BytesIO
from pypdf import PdfReader, PdfWriter
from docx import Document
from docx.shared import Inches

IMAGE_EXT = (".jpg", ".jpeg", ".png")

PDF_OUTPUT = "Consolidated_Documents.pdf"
DOC_OUTPUT = "Consolidated_Documents.docx"

writer = PdfWriter()
doc = Document()

files = sorted(os.listdir("."))

for file in files:

    # ---- HANDLE PDF FILES ----
    if file.lower().endswith(".pdf"):

        print("Processing PDF:", file)

        reader = PdfReader(file)

        # add pages to final PDF
        for page in reader.pages:
            writer.add_page(page)

        # convert each PDF page to image for DOCX
        pdf = fitz.open(file)

        for page in pdf:
            pix = page.get_pixmap(dpi=200)

            img_bytes = BytesIO(pix.tobytes("png"))

            doc.add_picture(img_bytes, width=Inches(6))
            doc.add_paragraph("")


    # ---- HANDLE IMAGE FILES ----
    elif file.lower().endswith(IMAGE_EXT):

        print("Processing Image:", file)

        img = Image.open(file)

        if img.mode != "RGB":
            img = img.convert("RGB")

        # add to PDF
        img_pdf_bytes = BytesIO()
        img.save(img_pdf_bytes, format="PDF")
        img_pdf_bytes.seek(0)

        img_pdf = PdfReader(img_pdf_bytes)
        writer.add_page(img_pdf.pages[0])

        # add to DOCX
        doc.add_picture(file, width=Inches(6))
        doc.add_paragraph("")


# save PDF
with open(PDF_OUTPUT, "wb") as f:
    writer.write(f)

# save DOCX
doc.save(DOC_OUTPUT)

print("Created:")
print(PDF_OUTPUT)
print(DOC_OUTPUT)