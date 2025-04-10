import os
import pandas as pd
import docx
from fpdf import FPDF
from openpyxl import load_workbook
import pdfplumber

# Function to convert Excel to PDF
def convert_excel_to_pdf(excel_path, output_filename="output.pdf"):
    df = pd.read_excel(excel_path)
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    for row in df.values:
        row_text = " | ".join(map(str, row))
        pdf.cell(200, 10, txt=row_text, ln=True)
    
    pdf.output(output_filename)
    print(f"PDF file '{output_filename}' saved.")

# Function to convert CSV to PDF
def convert_csv_to_pdf(csv_path, output_filename="output.pdf"):
    df = pd.read_csv(csv_path)
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    for row in df.values:
        row_text = " | ".join(map(str, row))
        pdf.cell(200, 10, txt=row_text, ln=True)
    
    pdf.output(output_filename)
    print(f"PDF file '{output_filename}' saved.")

# Function to convert Word to PDF
def convert_word_to_pdf(word_path, output_filename="output.pdf"):
    doc = docx.Document(word_path)
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    for para in doc.paragraphs:
        pdf.multi_cell(200, 10, txt=para.text)
    
    pdf.output(output_filename)
    print(f"PDF file '{output_filename}' saved.")

# Function to convert PDF to Excel
def convert_pdf_to_excel(pdf_path, output_filename="output.xlsx"):
    with pdfplumber.open(pdf_path) as pdf:
        all_text = []
        for page in pdf.pages:
            table = page.extract_table()
            if table:
                for row in table:
                    all_text.append(row)

        if all_text:
            # Convert the extracted table into a DataFrame
            df = pd.DataFrame(all_text[1:], columns=all_text[0])
            df.to_excel(output_filename, index=False)
            print(f"Excel file '{output_filename}' saved.")
        else:
            print("No table data found in the PDF.")

# Function to convert PDF to Word
def convert_pdf_to_word(pdf_path, output_filename="output.docx"):
    with pdfplumber.open(pdf_path) as pdf:
        doc = docx.Document()
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                doc.add_paragraph(text)
        
        doc.save(output_filename)
        print(f"Word file '{output_filename}' saved.")

# Function to convert Excel to Word
def convert_excel_to_word(excel_path, output_filename="output.docx"):
    df = pd.read_excel(excel_path)
    doc = docx.Document()

    for row in df.values:
        row_text = " | ".join(map(str, row))
        doc.add_paragraph(row_text)
    
    doc.save(output_filename)
    print(f"Word file '{output_filename}' saved.")

# Function to convert CSV to Word
def convert_csv_to_word(csv_path, output_filename="output.docx"):
    df = pd.read_csv(csv_path)
    doc = docx.Document()

    for row in df.values:
        row_text = " | ".join(map(str, row))
        doc.add_paragraph(row_text)
    
    doc.save(output_filename)
    print(f"Word file '{output_filename}' saved.")

# Function to convert Word to Excel
def convert_word_to_excel(word_path, output_filename="output.xlsx"):
    doc = docx.Document(word_path)
    data = []

    for para in doc.paragraphs:
        data.append([para.text])

    df = pd.DataFrame(data)
    df.to_excel(output_filename, index=False)
    print(f"Excel file '{output_filename}' saved.")

# Main function to ask the user for file input and output format
def main():
    print("Supported input formats: Excel, CSV, Word, PDF")
    input_file = input("Please provide the path to the input file: ").strip()

    # Check if the provided file path exists
    if not os.path.exists(input_file):
        print(f"The provided file '{input_file}' doesn't exist. Please check the path.")
        return

    # Get the file extension and determine output format
    file_extension = os.path.splitext(input_file)[1].lower()
    output_type = input("Please provide the output type (pdf, excel, word): ").strip().lower()

    # Convert based on input file extension and output type
    if file_extension == ".xlsx" or file_extension == ".xls":
        if output_type == "pdf":
            convert_excel_to_pdf(input_file)
        elif output_type == "word":
            convert_excel_to_word(input_file)
        elif output_type == "excel":
            print("Input file is already Excel format.")
        else:
            print("Unsupported conversion for Excel.")
    
    elif file_extension == ".csv":
        if output_type == "pdf":
            convert_csv_to_pdf(input_file)
        elif output_type == "word":
            convert_csv_to_word(input_file)
        elif output_type == "excel":
            print("Input file is already CSV format.")
        else:
            print("Unsupported conversion for CSV.")
    
    elif file_extension == ".docx":
        if output_type == "pdf":
            convert_word_to_pdf(input_file)
        elif output_type == "excel":
            convert_word_to_excel(input_file)
        elif output_type == "word":
            print("Input file is already Word format.")
        else:
            print("Unsupported conversion for Word.")
    
    elif file_extension == ".pdf":
        if output_type == "excel":
            convert_pdf_to_excel(input_file)
        elif output_type == "word":
            convert_pdf_to_word(input_file)
        elif output_type == "pdf":
            print("Input file is already PDF format.")
        else:
            print("Unsupported conversion for PDF.")
    
    else:
        print("Unsupported file type or conversion combination.")

if __name__ == "__main__":
    main()
